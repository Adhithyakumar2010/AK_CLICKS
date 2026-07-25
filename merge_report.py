from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

doc = Document('Photography_Story_Shop_Project_Report.docx')
doc.add_page_break()
doc.add_heading('Appendix A: Diagrams, Test Cases and Core Code', 1)
doc.add_paragraph('The following diagrams are provided in Mermaid notation. They can be rendered in GitHub, VS Code Markdown Preview, or Mermaid Live Editor.')

sections = [
('Use Case Diagram', '''flowchart LR\n  Visitor --> ViewPortfolio\n  Visitor --> SubmitBooking\n  Visitor --> BrowseStoryShop\n  Customer --> RegisterLogin\n  Customer --> PlaceOrder\n  Administrator --> ManageBookings\n  Administrator --> UpdateStock\n  Administrator --> ReviewNotifications'''),
('Sequence Diagram - Book Order', '''sequenceDiagram\n  Customer->>Browser: Add books to bag\n  Browser->>Flask: POST /order\n  Flask->>SQLite: Check stock and calculate total\n  SQLite-->>Flask: Availability result\n  Flask->>SQLite: Create order and reduce stock\n  Flask-->>Browser: Redirect to payment\n  Browser->>Flask: Confirm test payment\n  Flask->>SQLite: Save payment status and queue notice'''),
('Activity Diagram / Flowchart', '''flowchart TD\n  Start --> ChooseJourney\n  ChooseJourney -->|Photography| BookingForm\n  ChooseJourney -->|Story Shop| ShoppingBag\n  BookingForm --> SaveBooking\n  ShoppingBag --> CheckStock\n  CheckStock --> SaveOrder\n  SaveBooking --> Payment\n  SaveOrder --> Payment\n  Payment --> NotificationQueue\n  NotificationQueue --> AdminProcessing'''),
('System Architecture Diagram', '''flowchart TB\n  BrowserUI[Browser UI: HTML CSS JavaScript] --> Flask[Flask Application]\n  Flask --> Templates[Jinja Templates]\n  Flask --> SQLite[(SQLite Database)]\n  SQLite --> Bookings\n  SQLite --> Orders\n  SQLite --> ProductsStock\n  SQLite --> Customers\n  SQLite --> Notifications\n  Flask --> QR[QR Generator]\n  Flask --> Receipt[ReportLab PDF Receipts]'''),
('Workflow Chart', '''flowchart LR\n  Discover --> BookingOrOrder\n  BookingOrOrder --> Validation\n  Validation --> SQLiteStorage\n  SQLiteStorage --> TestPayment\n  TestPayment --> NotificationQueue\n  NotificationQueue --> AdminDashboard\n  AdminDashboard --> ApprovalOrFulfilment''')]
for title, code in sections:
    doc.add_heading(title, 2)
    p=doc.add_paragraph(); r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(8)

doc.add_heading('Real-Time Performance Considerations', 2)
for line in ['Portfolio and shop pages are served locally through Flask templates and static assets.', 'The availability endpoint runs a lightweight SQLite count query.', 'Order checkout validates stock and recalculates totals on the server.', 'QR codes and PDF receipts are generated on demand.', 'Production monitoring should measure page load, endpoint latency, database errors and notification delivery.']:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('Test Case Table', 2)
t=doc.add_table(rows=1, cols=3); t.style='Table Grid'
for cell,text in zip(t.rows[0].cells,['ID','Test case','Expected result']): cell.text=text
tests=[('TC-01','Submit valid booking','Booking record is saved and payment page opens.'),('TC-02','Check unused date','Availability API returns available true.'),('TC-03','Register customer','Password hash is stored and session starts.'),('TC-04','Order in-stock books','Order saves and stock decreases.'),('TC-05','Order unavailable stock','Order is rejected without stock change.'),('TC-06','Confirm payment method','Test payment state and method are stored.'),('TC-07','Download receipt','PDF receipt response is returned.'),('TC-08','Open admin without login','User is redirected to login.'),('TC-09','Update stock','Admin inventory amount is stored.'),('TC-10','Create order/booking','Notification queue receives a record.')]
for row in tests:
    cells=t.add_row().cells
    for c,v in zip(cells,row): c.text=v

doc.add_heading('Core Code Extracts', 2)
snippets=[('Booking creation route','''@app.route("/book", methods=["POST"])\ndef book():\n    fields = {key: request.form.get(key, "").strip() for key in fields}\n    with db_connection() as conn:\n        cursor = conn.execute("INSERT INTO bookings (...) VALUES (...)", fields)\n        booking_id = cursor.lastrowid\n    return redirect(url_for("payment", kind="booking", record_id=booking_id))'''),('Stock validation','''requested = Counter(item.strip() for item in customer["items"].split(",") if item.strip())\nall_products = {row["name"]: row for row in conn.execute("SELECT * FROM products")}\nfor name, quantity in requested.items():\n    conn.execute("UPDATE products SET stock=stock-? WHERE id=?", (quantity, all_products[name]["id"]))'''),('Availability API','''@app.route("/availability")\ndef availability():\n    date = request.args.get("date", "")\n    count = conn.execute("SELECT COUNT(*) FROM bookings WHERE booking_date=?", (date,)).fetchone()[0]\n    return jsonify({"date": date, "available": count == 0})''')]
for title,code in snippets:
    doc.add_heading(title,3); p=doc.add_paragraph(); r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(8)
doc.save('AK_CLICKS_Complete_Project_Report.docx')
