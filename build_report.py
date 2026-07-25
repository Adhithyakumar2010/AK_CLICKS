from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = 'Photography_Story_Shop_Project_Report.docx'
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def set_cell(cell, text, bold=False, color=None):
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(text)); r.bold=bold; r.font.name='Calibri'; r.font.size=Pt(9); 
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): shade(t.rows[0].cells[i], 'F2F4F7'); set_cell(t.rows[0].cells[i],h,True,'1F4D78')
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row): set_cell(cells[i],val)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t
def p(text='', style=None, bold=False, italic=False, align=None, size=None, color=None):
    para=doc.add_paragraph(style=style); para.paragraph_format.space_after=Pt(6); para.paragraph_format.line_spacing=1.1
    if align: para.alignment=align
    r=para.add_run(text); r.bold=bold; r.italic=italic; r.font.name='Calibri'
    if size:r.font.size=Pt(size)
    if color:r.font.color.rgb=RGBColor.from_string(color)
    return para
def h(text, level=1):
    para=doc.add_heading(text, level=level); para.paragraph_format.space_before=Pt(16 if level==1 else 12); para.paragraph_format.space_after=Pt(7)
    for r in para.runs: r.font.name='Calibri'; r.font.color.rgb=RGBColor(46,116,181) if level<3 else RGBColor(31,77,120)
def bullets(items):
    for x in items: p(x, 'List Bullet')

# cover
p('PROJECT REPORT', bold=True, size=11, color='7A5A00', align=WD_ALIGN_PARAGRAPH.CENTER)
p('AK CLICKS: Photography Booking and The Story Shop Platform', bold=True, size=28, color='0B2545', align=WD_ALIGN_PARAGRAPH.CENTER)
p('A combined Flask web application for event photography bookings, book commerce, customer accounts and studio operations', italic=True, size=14, color='555555', align=WD_ALIGN_PARAGRAPH.CENTER)
p('', size=20)
table(['Prepared for','Technology','Project type'], [['AK CLICKS','Python Flask, SQLite, HTML/CSS/JavaScript','Full-stack portfolio, booking and commerce website']], [2.1,2.2,2.2])
p('July 2026', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

h('Executive Summary')
p('AK CLICKS is a full-stack web platform that combines a photography portfolio and event booking system with The Story Shop, an online bookstore. Customers can explore photography services, submit wedding, birthday, event, portrait or brand-content enquiries, purchase books, create accounts and download test payment receipts. Studio administrators use one dashboard to manage bookings, orders, inventory, payment submissions and queued notifications.')
h('Project Title',2); p('AK CLICKS - Integrated Photography Booking and Story Shop Management System', bold=True)
h('Objectives',2); bullets(['Present photography work in a professional responsive portfolio.','Capture online enquiries for weddings, birthdays and other events.','Sell Story Shop books through a stock-aware shopping bag and checkout flow.','Provide a shared admin workspace for bookings, orders, stock, payment state and notifications.','Create a foundation for live payments, email and WhatsApp delivery.'])
h('Problem Statement')
p('Photography businesses often handle enquiries through calls and messages, while independent shops track orders and inventory separately. This creates missed dates, manual stock errors and inconsistent customer communication. The project centralises these workflows in a single browser-based application.')

h('Scope and User Roles')
table(['Role','Key capabilities'], [['Visitor','View photography portfolio, browse books, check a date and submit booking/order details.'],['Customer','Register, sign in and view their Story Shop order history.'],['Administrator','Approve or decline bookings; process orders; update book stock; review payment state and notification queue.']], [1.35,5.15])

h('Technology Stack')
table(['Layer','Technology','Purpose'], [['Frontend','HTML5, CSS3, JavaScript','Responsive interface, bag management, form controls and payment-method switching.'],['Backend','Python 3 with Flask','Routes, validation, sessions, payment pages, PDF receipt generation and admin access.'],['Database','SQLite','Persistent bookings, orders, customers, products, stock and notifications.'],['Security utilities','Werkzeug password hashing','Stores customer passwords as hashes rather than plain text.'],['QR and receipts','qrcode, ReportLab','Local SVG UPI QR generation and downloadable PDF test receipts.']], [1.25,1.8,3.45])

h('System Architecture')
p('The browser renders Jinja HTML templates served by Flask. Forms submit to Flask routes, which validate data and write to SQLite. The same data is read by the customer account and administrator dashboard. JavaScript keeps the temporary shopping bag in browser local storage until checkout. The architecture is intentionally simple for local deployment and can later be moved to PostgreSQL and a managed hosting service.')
table(['Component','Responsibility','Primary data'], [['Public site','Portfolio, booking form, availability check','Bookings'],['Story Shop','Catalog, bag, checkout','Products, orders, stock'],['Payment module','Test payment method, QR, receipt','Payment method/status'],['Account module','Registration, login, order history','Customers, orders'],['Admin dashboard','Operations control','All operational records']], [1.25,2.45,2.8])

h('Functional Modules')
h('Photography Booking Module',2)
bullets(['Services include wedding, birthday, event, portrait and brand photography.','The selected event date can be checked through the availability endpoint. A date with a non-declined booking is reported as already having an enquiry.','A valid enquiry stores customer contact details, service type, event date, location and notes, then opens the booking-advance payment page.'])
h('Story Shop Module',2)
bullets(['The integrated catalog contains six books and uses the original Story Shop artwork.','The bag calculates the total in the browser; the server recalculates the final total and checks stock before creating an order.','Stock is decremented only after a valid order is created, preventing orders for unavailable quantities.'])
h('Customer Account Module',2)
bullets(['New customers can register with name, email and a minimum six-character password.','Passwords are protected with Werkzeug password hashes.','Signed-in customers can review orders associated with their account.'])
h('Payment and Receipt Module',2)
p('The payment screens provide UPI/QR, net-banking and card user interfaces. They are intentionally labelled TEST MODE: no financial data is transferred and no real money is collected. After a customer selects a method, the system records a test payment submission and creates a PDF receipt that explicitly states it is not proof of a financial transaction.')

h('Database Design')
table(['Table','Purpose','Important fields'], [['bookings','Photography enquiries','name, email, phone, service, booking_date, status, payment_status'],['orders','Story Shop purchases','customer_name, address, items, total, status, customer_id'],['customers','Customer identities','name, email (unique), password_hash'],['products','Book catalog and inventory','id, name, price, stock, image, description'],['notifications','Queued communications','recipient, channel, subject, body, status']], [1.25,1.55,3.7])

h('Routes and API Reference')
table(['Method','Route','Purpose'], [['GET','/','Photography portfolio and booking form'],['POST','/book','Creates booking enquiry and opens booking payment'],['GET','/availability?date=YYYY-MM-DD','JSON date-availability response'],['GET','/shop','Story Shop catalog and bag'],['POST','/order','Validates stock, creates order and reserves inventory'],['GET/POST','/account','Customer registration and login'],['GET','/payment/<kind>/<id>','Payment-selection page'],['GET','/payment/<kind>/<id>/qr','SVG test-mode UPI QR code'],['POST','/payment/<kind>/<id>/confirm','Stores test payment submission'],['GET','/receipt/<kind>/<id>.pdf','Downloads a test payment receipt'],['GET','/admin','Protected studio dashboard']], [0.65,2.45,3.4])

h('Workflow: Booking to Confirmation')
for step in ['Customer completes the booking form.','Flask validates required fields and writes a bookings record.','The system queues a confirmation notification record.','Customer selects a test payment method for the configured booking advance.','The status and method are saved; a receipt can be downloaded.','Administrator reviews the enquiry and approves, declines or leaves it pending.']: p(step, 'List Number')
h('Workflow: Order to Fulfilment')
for step in ['Customer adds books to the browser bag.','Checkout submits contact and delivery details.','Server checks each product against current stock, recalculates the total and decrements inventory.','Order and queued confirmation record are created.','Customer performs a test payment submission and downloads a receipt.','Administrator changes order status to New, Processing or Fulfilled.']: p(step, 'List Number')

h('Notification Design')
p('Email and WhatsApp notifications are implemented as a local queue. Booking, order and test payment actions insert notification records visible in the admin dashboard. This approach proves the business workflow without sending external messages or requiring credentials. A production integration should use SMTP or a transactional email provider plus WhatsApp Business Cloud API or Twilio, using credentials stored only as environment variables.')

h('Security and Validation')
bullets(['Administrator dashboard is protected by a session login; production credentials must replace the current development defaults.','Customer passwords are hashed with Werkzeug.','Checkout recalculates price and stock on the server rather than trusting the client total.','Payment forms are explicitly test mode and must never collect actual card details until a PCI-compliant gateway is integrated.','Production deployment must use HTTPS, a strong SECRET_KEY, CSRF protection, secure cookies and a production database.'])

h('Testing Performed')
table(['Area','Verification'], [['Application syntax','Python compilation completed successfully.'],['Customer account','Registration and account-page access tested using an isolated in-memory database.'],['Shop and stock-aware order','Catalog response, valid order creation and inventory validation tested.'],['Availability API','JSON response checked for an unused date.'],['Payment and QR','Payment page, confirmation route and SVG QR endpoint checked.'],['Receipt','PDF receipt response generated and confirmed as application/pdf.']], [1.9,4.6])

h('Limitations and Production Roadmap')
table(['Current state','Production next step'], [['Test-mode payment UI and QR','Integrate Razorpay/Stripe and use verified webhooks before marking payments as paid.'],['Queued notifications','Configure SMTP/SendGrid and WhatsApp Business/Twilio credentials.'],['SQLite local database','Use PostgreSQL with backups and migration tooling.'],['Simple date availability','Add time slots, staff/resource capacity, calendar UI and calendar sync.'],['Local server','Deploy behind HTTPS on a managed host with environment secrets.'],['Basic customer area','Add password reset, verified email, booking history and profile editing.']], [2.65,3.85])

h('Conclusion')
p('The AK CLICKS platform demonstrates a practical integrated web application for a photography business and an independent book store. It turns portfolio visits into trackable booking enquiries, turns book selections into stock-aware orders, and gives studio staff one operational view. The project is suitable for demonstration and local use today, with a clear path to live payments, external communications and production deployment.')

# footer
for section in doc.sections:
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=fp.add_run('AK CLICKS Project Report | July 2026'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
doc.save(OUT)
print(OUT)
